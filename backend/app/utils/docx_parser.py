import logging
from docx import Document
from fastapi import HTTPException

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


class DOCXParser:
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"

            if not text.strip():
                raise HTTPException(
                    status_code=400, detail="No text could be extracted from the DOCX")

            logger.debug(f"Extracted text from DOCX: {text[:500]}...")

            return text.strip()
        except Exception as e:
            logger.error(
                f"Error parsing DOCX {file_path}: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=500, detail=f"Error parsing DOCX: {str(e)}")
